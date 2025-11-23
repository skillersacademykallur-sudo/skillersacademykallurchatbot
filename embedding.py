import os
import time
import shutil
import json
import logging
import pandas as pd

from PyPDF2 import PdfReader
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain.docstore.document import Document as LangChainDocument

from pinecone import Pinecone, ServerlessSpec
from dotenv import load_dotenv
from huggingface_hub import snapshot_download


# -------------------- LOGGING --------------------
LOG_DIR = r"C:\Storage\DS\Projects\SkillersChatbotlogfiles\EmbeddingLogs"
os.makedirs(LOG_DIR, exist_ok=True)

logging.basicConfig(
    filename=os.path.join(LOG_DIR, "embedding.log"),
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

logging.info("===== Embedding Job Started =====")


# -------------------- CONFIG --------------------
BASE_DIR = r"C:\Storage\DS\Projects\skillersacademykallurchatbot"
DATA_FOLDER = os.path.join(BASE_DIR, "Data")
ENV_PATH = os.path.join(BASE_DIR, ".env")

LOCAL_MODEL_PATH = os.path.join(BASE_DIR, "models", "all-MiniLM-L12-v2")

load_dotenv(ENV_PATH)

PINECONE_API_KEY = os.getenv('PINECONE_API_KEY')
INDEX_NAME = "skillersbot12112025"

logging.info(f"Data directory: {DATA_FOLDER}")


# -------------------- JSON → CLEAN TEXT --------------------
def json_to_clean_text(obj, prefix=""):
    """
    Convert JSON into natural readable text (not dict format).
    Example:
      name: Tikri IT Services
      phone: 1234
    """
    lines = []

    if isinstance(obj, dict):
        for key, value in obj.items():
            new_prefix = f"{key}: "
            lines.append(json_to_clean_text(value, new_prefix))

    elif isinstance(obj, list):
        for i, item in enumerate(obj, 1):
            new_prefix = f"{prefix}Item {i} - "
            lines.append(json_to_clean_text(item, new_prefix))

    else:
        lines.append(f"{prefix}{obj}")

    return "\n".join(lines)


# -------------------- LOAD EMBEDDING MODEL --------------------
def get_embeddings_model():
    model_repo = "sentence-transformers/all-MiniLM-L12-v2"

    try:
        if not os.path.exists(LOCAL_MODEL_PATH) or len(os.listdir(LOCAL_MODEL_PATH)) < 5:
            logging.warning("Local model missing. Downloading...")

            if os.path.exists(LOCAL_MODEL_PATH):
                shutil.rmtree(LOCAL_MODEL_PATH)

            snapshot_download(
                repo_id=model_repo,
                local_dir=LOCAL_MODEL_PATH,
                local_dir_use_symlinks=False,
                ignore_patterns=["*.h5", "*.ot", "*.onnx", "*.msgpack", "rust_model.ot"]
            )

            logging.info("Model downloaded successfully.")

        return HuggingFaceEmbeddings(
            model_name=LOCAL_MODEL_PATH,
            model_kwargs={"device": "cpu"}
        )

    except Exception as e:
        logging.error(f"Error loading embeddings model: {e}", exc_info=True)
        raise e


# -------------------- EXTRACT TEXT FROM EACH FILE --------------------
def extract_text_from_file(file_path, filename):
    text = ""

    try:
        # ---------- JSON ----------
        if filename.lower().endswith(".json"):
            logging.info(f"Processing JSON: {filename}")
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            text = json_to_clean_text(data)

        # ---------- TXT ----------
        elif filename.lower().endswith(".txt"):
            logging.info(f"Processing TXT: {filename}")
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        # ---------- DOCX ----------
        elif filename.lower().endswith(".docx"):
            logging.info(f"Processing DOCX: {filename}")
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])

        # ---------- PDF ----------
        elif filename.lower().endswith(".pdf"):
            logging.info(f"Processing PDF: {filename}")
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)

        # ---------- EXCEL ----------
        elif filename.lower().endswith((".xlsx", ".xls")):
            logging.info(f"Processing Excel: {filename}")
            excel_data = pd.read_excel(file_path, sheet_name=None)
            txt = []
            for sheet_name, df in excel_data.items():
                txt.append(f"Sheet: {sheet_name}\n")
                txt.append(df.to_string(index=False))
            text = "\n".join(txt)

        # ---------- IMAGES (OCR) ----------
        elif filename.lower().endswith((".png", ".jpg", ".jpeg")):
            logging.info(f"Processing Image: {filename}")
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)

    except Exception as e:
        logging.error(f"Error processing {filename}: {e}", exc_info=True)

    return text.strip()


# -------------------- CREATE DOCUMENTS FOR EMBEDDING --------------------
def build_documents():
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
    documents = []

    for filename in os.listdir(DATA_FOLDER):
        file_path = os.path.join(DATA_FOLDER, filename)
        text = extract_text_from_file(file_path, filename)

        if not text:
            continue

        # If text is small (like your JSON records) → do NOT split
        if len(text) < 1500:
            documents.append(
                LangChainDocument(page_content=text, metadata={"source": filename})
            )
        else:
            chunks = splitter.split_text(text)
            for ch in chunks:
                documents.append(
                    LangChainDocument(page_content=ch, metadata={"source": filename})
                )

    logging.info(f"Total documents prepared: {len(documents)}")
    return documents


# -------------------- MAIN: EMBED + UPSERT TO PINECONE --------------------
def run_embedding_job():
    try:
        embeddings = get_embeddings_model()
        documents = build_documents()

        pc = Pinecone(api_key=PINECONE_API_KEY)

        # Create index if not exists
        existing = [i.name for i in pc.list_indexes()]
        if INDEX_NAME not in existing:
            logging.info(f"Creating Pinecone index: {INDEX_NAME}")
            pc.create_index(
                name=INDEX_NAME,
                dimension=384,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )

        logging.info("Upserting documents into Pinecone...")

        PineconeVectorStore.from_documents(
            documents=documents,
            index_name=INDEX_NAME,
            embedding=embeddings
        )

        logging.info("Embedding + Upsert completed successfully.")

    except Exception as e:
        logging.error(f"Failed embedding job: {e}", exc_info=True)


if __name__ == "__main__":
    run_embedding_job()
    logging.info("===== Embedding Job Completed =====")
