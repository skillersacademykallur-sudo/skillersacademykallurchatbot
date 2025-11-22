import os
import time
import shutil  # Added for cleaning up corrupt folders
import pandas as pd
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain.docstore.document import Document as LangChainDocument
from pinecone import Pinecone, ServerlessSpec
from dotenv import load_dotenv
from huggingface_hub import snapshot_download

# --- 1. SETUP & CONFIGURATION ---
# Define absolute paths (Crucial for automation)
BASE_DIR = r"C:\Storage\DS\Projects\skillersacademykallurchatbot"
DATA_FOLDER = os.path.join(BASE_DIR, "Data")
ENV_PATH = os.path.join(BASE_DIR, ".env")
# Define a permanent folder for the model to sit in
LOCAL_MODEL_PATH = os.path.join(BASE_DIR, "models", "all-MiniLM-L12-v2")

# Load Environment Variables
load_dotenv(ENV_PATH)

PINECONE_API_KEY = os.getenv('PINECONE_API_KEY')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
INDEX_NAME = "skillersbot12112025"

print(f"Starting Daily Update at {time.strftime('%Y-%m-%d %H:%M:%S')}")


# --- 2. ROBUST & OPTIMIZED MODEL LOADING FUNCTION ---
def get_embeddings_model():
    """
    Checks if model exists locally and is complete.
    If yes, load it.
    If no, clean up, download ONLY the necessary files, then load it.
    """
    model_repo = "sentence-transformers/all-MiniLM-L12-v2"

    # Check if the folder exists AND has enough files (heuristic to detect incomplete download)
    # A complete model usually has > 5 files.
    if not os.path.exists(LOCAL_MODEL_PATH) or len(os.listdir(LOCAL_MODEL_PATH)) < 5:
        print(f"⚠️ Model missing or incomplete at {LOCAL_MODEL_PATH}")

        # Clean up potential corrupt/partial folder before downloading
        if os.path.exists(LOCAL_MODEL_PATH):
            print("🧹 Removing partial/corrupt model folder...")
            shutil.rmtree(LOCAL_MODEL_PATH)

        print(f"⬇️ Downloading optimized version of {model_repo}...")
        try:
            snapshot_download(
                repo_id=model_repo,
                local_dir=LOCAL_MODEL_PATH,
                local_dir_use_symlinks=False,
                # THIS LINE SAVES TIME: Ignores heavy files you don't need (TensorFlow, Rust, ONNX)
                ignore_patterns=["*.h5", "*.ot", "*.onnx", "*.msgpack", "rust_model.ot", "tf_model.h5"]
            )
            print("✅ Download complete.")
        except Exception as e:
            print(f"❌ Critical Error downloading model: {e}")
            raise e
    else:
        print(f"✅ Local model found at {LOCAL_MODEL_PATH}")

    # Load the model from the local folder
    print("🔌 Loading model into memory...")
    return HuggingFaceEmbeddings(
        model_name=LOCAL_MODEL_PATH,
        model_kwargs={"device": "cpu"}  # Change to 'cuda' if you have a GPU
    )


# --- 3. TEXT EXTRACTION ---
all_text = ""

if not os.path.exists(DATA_FOLDER):
    print(f"Error: Folder '{DATA_FOLDER}' not found!")
else:
    for filename in os.listdir(DATA_FOLDER):
        file_path = os.path.join(DATA_FOLDER, filename)
        text = ""

        try:
            # PDF
            if filename.lower().endswith(".pdf"):
                print(f"Processing PDF: {filename}")
                reader = PdfReader(file_path)
                text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                if not text.strip():
                    print(f"Performing OCR on {filename}...")
                    images = convert_from_path(file_path)
                    text = "\n".join(pytesseract.image_to_string(img) for img in images)

            # Word
            elif filename.lower().endswith(".docx"):
                print(f"Processing Word: {filename}")
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])

            # TXT
            elif filename.lower().endswith(".txt"):
                if filename == "all_combined_text.txt": continue
                print(f"Processing TXT: {filename}")
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()

            # Excel
            elif filename.lower().endswith((".xlsx", ".xls")):
                print(f"Processing Excel: {filename}")
                excel_data = pd.read_excel(file_path, sheet_name=None)
                text_parts = []
                for sheet_name, df in excel_data.items():
                    text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                    text_parts.append(df.to_string(index=False))
                text = "\n".join(text_parts)

            if text.strip():
                all_text += f"\n\n=== FILE: {filename} ===\n{text}"

        except Exception as e:
            print(f"Error processing {filename}: {e}")

# Save combined text
combined_file_path = os.path.join(DATA_FOLDER, "all_combined_text.txt")
with open(combined_file_path, "w", encoding="utf-8") as f:
    f.write(all_text)
print("✅ All text combined.")

# --- 4. CHUNKING ---
if not all_text:
    print("No text found. Exiting.")
    exit()

text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=20)
text_chunks = text_splitter.split_text(all_text)
print(f"✅ Total chunks created: {len(text_chunks)}")

# --- 5. EMBEDDING & PINECONE UPSERT ---
# Initialize Embeddings using our Robust Function
embeddings = get_embeddings_model()

# Initialize Pinecone
pc = Pinecone(api_key=PINECONE_API_KEY)

# Check/Create Index
existing_indexes = [index.name for index in pc.list_indexes()]
if INDEX_NAME not in existing_indexes:
    print(f"Creating index: {INDEX_NAME}")
    pc.create_index(
        name=INDEX_NAME,
        dimension=384,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )

# Prepare documents
documents = [LangChainDocument(page_content=chunk) for chunk in text_chunks]

# Upsert
print("📤 Upserting to Pinecone...")
PineconeVectorStore.from_documents(
    documents=documents,
    index_name=INDEX_NAME,
    embedding=embeddings
)

print("✅ Pinecone Index Updated Successfully.")
print("Job Finished.")